import openai
import streamlit as st
import os
from streamlit_chat import message as msg
import docx
import io

openai.__version__
'0.27.8'

openai.api_key = os.getenv = "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"

st.title("Hy2nIA Chat com ChatGPT Turbo")

st.subheader("Chat")

if 'hst_conversa' not in st.session_state:
    st.session_state.hst_conversa = []

conversa = st.text_area("Digite a pegurta:")
btn_enviar_msg = st.button("Enviar Mensagem")
if btn_enviar_msg:
    st.session_state.hst_conversa.append({"role": "user", "content": conversa})
    retorno_openai = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=st.session_state.hst_conversa,
        max_tokens=500,
        n=1
    )
    st.session_state.hst_conversa.append({"role": "assistant",
                                          "content": retorno_openai['choices'][0]['message']['content']})

if len(st.session_state.hst_conversa) > 0:
    for i in range(len(st.session_state.hst_conversa)):
        if i % 2 == 0:
            msg("Você: " + st.session_state.hst_conversa[i]['content'], is_user=True)
        else:
            msg("Reposta IA: " + st.session_state.hst_conversa[i]['content'])

if len(st.session_state.hst_conversa) > 0:
    btn_salvar = st.button("Salvar Chat")
    if btn_salvar:
        trabalho = io.BytesIO()
        documento = docx.Document()
        documento.add_heading('Conteúdo Gerado', level=1)

        for i in range(len(st.session_state.hst_conversa)):
            if i % 2 == 0:
                documento.add_heading("Pergunta", level=2)
                documento.add_paragraph(st.session_state.hst_conversa[i]['content'])
            else:
                documento.add_heading("Resposta", level=2)
                documento.add_paragraph(st.session_state.hst_conversa[i]['content'])

            documento.save(trabalho)
            st.download_button(label="Clique aqui para fazer o download",
                               data=trabalho,
                               file_name="Hy2nIA - Chat com o Especialista.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
